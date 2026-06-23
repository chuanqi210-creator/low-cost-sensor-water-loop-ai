from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().with_name("研究方案_Word兼容版.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 35, 50)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F4F6F9"
LIGHT_BLUE = "E8EEF5"
BORDER = "BFC7D2"

LATIN_FONT = "Arial"
EAST_ASIA_FONT = "SimSun"
BODY_FONT = LATIN_FONT
HEADING_FONT = LATIN_FONT


def set_east_asia_font(run, font_name):
    run.font.name = LATIN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rfonts.set(qn("w:hint"), "eastAsia")


def set_run(run, size=None, color=None, bold=None, italic=None, font=BODY_FONT):
    set_east_asia_font(run, font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_p_spacing(p, before=0, after=6, line=1.25, align=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align


def patch_style(style, font=BODY_FONT, size=11, color=INK, before=0, after=6, line=1.25, bold=False):
    style.font.name = LATIN_FONT
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rfonts.set(qn("w:hint"), "eastAsia")
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def patch_doc_defaults(doc):
    """Force Word-compatible Chinese font slots at document-default level."""
    styles_el = doc.styles._element
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rfonts.set(qn("w:hint"), "eastAsia")

    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_p_spacing(p, before=before, after=after, line=line, align=align)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_mixed_para(doc, parts, before=0, after=6, line=1.25, align=None):
    p = doc.add_paragraph()
    set_p_spacing(p, before=before, after=after, line=line, align=align)
    for part in parts:
        text = part.get("text", "")
        r = p.add_run(text)
        set_run(
            r,
            size=part.get("size", 11),
            color=part.get("color", INK),
            bold=part.get("bold", False),
            italic=part.get("italic", False),
            font=part.get("font", BODY_FONT),
        )
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        set_p_spacing(p, before=16, after=8, line=1.25)
    elif level == 2:
        set_p_spacing(p, before=12, after=6, line=1.25)
    else:
        set_p_spacing(p, before=8, after=4, line=1.25)
    for r in p.runs:
        set_run(r, font=HEADING_FONT)
    return p


def add_label_block(doc, label, text):
    p = doc.add_paragraph()
    set_p_spacing(p, before=0, after=6, line=1.25)
    r1 = p.add_run(label)
    set_run(r1, size=11, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=11, color=INK)
    return p


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def write_cell(cell, text, bold=False, color=INK, size=10.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_p_spacing(p, before=0, after=0, line=1.18, align=align)
    r = p.add_run(text)
    set_run(r, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_FILL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        write_cell(hdr[i], h, bold=True, color=DARK_BLUE, size=10.2, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            write_cell(cells[i], val, size=9.8)
    add_para(doc, "", after=2)
    return table


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="D5DAE2", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    cell.text = ""
    p = cell.paragraphs[0]
    set_p_spacing(p, before=0, after=0, line=1.25)
    r1 = p.add_run(title)
    set_run(r1, size=11, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run(r2, size=11, color=INK)
    add_para(doc, "", after=2)


def section_break(doc):
    # Let Word paginate naturally; explicit breaks can render as visible
    # placeholder boxes in some Quick Look previews.
    return None


doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)
    s.header_distance = Inches(0.492)
    s.footer_distance = Inches(0.492)

patch_style(doc.styles["Normal"], font=BODY_FONT, size=11, color=INK, before=0, after=6, line=1.25)
patch_style(doc.styles["Heading 1"], font=HEADING_FONT, size=16, color=BLUE, before=16, after=8, line=1.25, bold=True)
patch_style(doc.styles["Heading 2"], font=HEADING_FONT, size=13, color=BLUE, before=12, after=6, line=1.25, bold=True)
patch_style(doc.styles["Heading 3"], font=HEADING_FONT, size=12, color=DARK_BLUE, before=8, after=4, line=1.25, bold=True)
patch_doc_defaults(doc)

# Running footer
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("低成本传感循环式水处理重大项目研究方案")
set_run(run, size=9, color=MUTED)

# Cover
add_para(doc, "重大项目级研究方案总括版", size=12, color=MUTED, bold=True, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    doc,
    "低成本传感条件下循环式水处理结构、软传感器灰箱推断、多智能体机理诊断与闭环控制研究方案",
    size=22,
    color=RGBColor(0, 0, 0),
    bold=True,
    before=6,
    after=8,
    line=1.2,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_para(
    doc,
    "面向复杂废水处理过程黑箱问题的灰箱感知与智能闭环控制框架",
    size=13,
    color=MUTED,
    after=20,
    align=WD_ALIGN_PARAGRAPH.CENTER,
)
add_para(doc, "整理日期：2026 年 5 月 31 日", size=10.5, color=MUTED, after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "版本：讨论汇总与研究规划版", size=10.5, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)

add_callout(
    doc,
    "核心命题：",
    "在低成本、慢速、不完整传感条件下，通过循环式水处理结构把时间转化为信息，利用软传感器估计不可直接观测的过程状态，再由多智能体系统完成机理解释、故障诊断与闭环控制，动态决定是否回流、延长停留时间、调整药剂投加或进入下一处理单元。",
)

add_table(
    doc,
    ["模块", "核心作用", "对应问题"],
    [
        ["循环式结构", "提供回流、暂存、延长停留和风险拦截能力", "传感慢、反应慢、判断慢"],
        ["低成本传感", "以 pH、ORP、电导率、浊度、温度、流量等形成稀疏证据链", "昂贵全量传感不可负担"],
        ["软传感器", "估计污染物残留、反应进程、氧化剂余量、催化剂活性等隐藏状态", "中间过程不可直接观测"],
        ["多智能体系统", "完成数据质控、机理解释、故障诊断、成本安全评价和决策仲裁", "单一黑箱模型缺乏解释和行动能力"],
        ["闭环控制", "执行回流、延时、加药、调 pH、切换单元、报警或放行", "预测结果无法自然转化为工程动作"],
    ],
    [1700, 4100, 3560],
)

section_break(doc)

add_heading(doc, "一、项目摘要", 1)
add_para(
    doc,
    "复杂废水处理过程常处于低可观测状态：在线高精度传感器价格高、维护难，实际工程中往往只能获得 pH、ORP、电导率、浊度、温度、流量等低成本信号；而污染物降解、氧化剂消耗、催化剂失活、基质干扰、传质限制等关键过程状态并不能被直接观测。传统机器学习方法多将问题简化为输入到出水指标的黑箱预测，难以解释异常原因，也难以稳定转化为可执行控制动作。",
)
add_para(
    doc,
    "本方案提出一种循环式灰箱感知与多智能体闭环控制框架。其基本思想是：不依赖昂贵全量传感器硬拆黑箱，而是通过循环、回流、暂存和延长停留时间创造可观测时间窗口；利用软传感器从低成本信号和循环历史中估计不可直接观测状态；再由多智能体系统进行机理解释、故障诊断、控制策略生成、成本安全评价和仲裁；最终驱动泵、阀、加药、pH 调节、回流和单元切换等工程动作。",
)
add_para(
    doc,
    "该方案的目标不是一次性建立一个全自动水厂，而是形成一套可扩展的研究范式：在低成本传感条件下，把“进水可见、出水可见、中间黑箱”的传统过程，改造为“稀疏可感知、状态可估计、机理可解释、动作可闭环”的灰箱智能水处理系统。",
)

add_heading(doc, "二、立项依据与问题定义", 1)
add_heading(doc, "2.1 工程痛点", 2)
add_label_block(doc, "痛点一：", "传感器成本与维护约束导致过程不可见。高精度在线检测设备价格高、维护复杂，在复杂工业废水、含氟废水、PFAS/PFOA 或高级氧化体系中，很难做到全量、高频、长期稳定监测。")
add_label_block(doc, "痛点二：", "黑箱机器学习难以形成控制行动。单纯预测出水浓度并不等于知道系统为什么异常，更不等于知道应该回流、加药、调 pH 还是切换处理单元。")
add_label_block(doc, "痛点三：", "一次通过式处理缺少诊断时间。若传感和反应速度不够快，处理系统来不及判断过程是否完成，也来不及根据异常采取纠偏动作。")
add_label_block(doc, "痛点四：", "传统回流缺少智能决策。固定回流可以提高处理机会，但也可能带来无效循环、能耗上升、药剂浪费和处理能力下降。")

add_heading(doc, "2.2 本项目的基本判断", 2)
add_callout(
    doc,
    "关键判断：",
    "低成本传感本身不是缺陷，而是必须正视的工程约束；循环结构也不是简单重复处理，而是为慢速传感、灰箱估计、多智能体诊断和闭环控制争取时间的信息结构。",
)
add_para(
    doc,
    "因此，本项目关注的不是“如何购买更多昂贵传感器”，而是“如何在有限传感条件下设计一套能够逐步看清过程、解释异常并采取行动的系统”。这也是本方案区别于传统水质预测、传统回流工艺和单一自动控制系统的核心。",
)

add_heading(doc, "三、总体目标与核心科学问题", 1)
add_heading(doc, "3.1 总体目标", 2)
add_para(
    doc,
    "构建面向复杂废水处理过程的循环式灰箱智能控制系统，在低成本传感条件下，通过循环结构争取诊断与反应时间，利用软传感器估计不可直接观测状态，再由多智能体系统完成机理解释、故障诊断与闭环控制，实现对回流、停留时间、加药、pH 调节、单元切换和风险放行的动态决策。",
)

add_heading(doc, "3.2 核心科学问题", 2)
add_table(
    doc,
    ["编号", "科学问题", "需要突破的关键点"],
    [
        ["SQ1", "低成本稀疏传感下如何恢复隐藏过程状态？", "从少量、噪声、延迟传感信号中估计污染物残留、反应进程、氧化剂余量和催化剂活性。"],
        ["SQ2", "循环结构如何从工艺回流升级为信息增益机制？", "明确回流何时有价值，何时应停止循环，如何评价回流边际收益。"],
        ["SQ3", "软传感器如何从出水预测升级为灰箱状态估计？", "输出过程状态向量、不确定性和达标概率，而不是单一浓度预测。"],
        ["SQ4", "多智能体如何完成证据约束的诊断与仲裁？", "让不同 agent 独立承担数据质控、机理解释、故障诊断、控制建议和安全成本评价。"],
        ["SQ5", "如何把诊断结果转化为可执行闭环动作？", "动态选择回流、延长停留、补加药剂、调 pH、切换单元、暂存待检或报警。"],
    ],
    [900, 3350, 5110],
)

add_heading(doc, "3.3 研究假设", 2)
add_label_block(doc, "H1：", "低成本传感虽然单点信息不足，但在循环结构中形成多轮时间序列后，可以显著提高隐藏状态估计能力。")
add_label_block(doc, "H2：", "循环停留能够为慢速传感、慢反应和模型诊断争取时间，从而降低对昂贵高频传感器的依赖。")
add_label_block(doc, "H3：", "多智能体诊断系统比单一黑箱模型更适合处理原因解释、故障识别和安全控制建议。")
add_label_block(doc, "H4：", "引入回流边际收益后，系统可以避免无效循环，在达标率、成本、能耗、时间和安全之间取得更优平衡。")

add_heading(doc, "四、总体技术路线", 1)
add_para(
    doc,
    "项目技术路线可概括为：循环式水处理结构提供可执行的物理空间；低成本传感器采集过程信号；软传感器将稀疏信号转化为隐藏状态估计；多智能体系统对状态进行机理解释和故障诊断；闭环控制器将诊断结果转化为工程动作；执行后的新状态再进入下一轮感知与判断。",
)
add_callout(
    doc,
    "技术路线：",
    "进水缓冲与水质扰动识别 -> 主反应器处理 -> 低成本在线传感 -> 软传感器灰箱状态估计 -> 多智能体机理诊断与风险仲裁 -> 控制动作执行 -> 回流、延长停留、补加药剂、调 pH、切换单元或达标放行 -> 新一轮反馈。",
)
add_table(
    doc,
    ["层级", "输入", "核心方法", "输出"],
    [
        ["物理过程层", "进水、水质负荷、反应器状态", "循环、回流、暂存、加药、单元切换", "可调节水处理过程"],
        ["传感层", "pH、ORP、电导率、浊度、温度、流量、UV254 等", "传感器校准、漂移识别、数据清洗", "稀疏过程证据链"],
        ["灰箱估计层", "传感信号、历史循环、投加记录、离线真值", "软传感器、时序学习、机理约束、不确定性估计", "隐藏状态向量"],
        ["多智能体层", "隐藏状态、知识图谱、故障库、成本安全约束", "多 agent 盲评、一致性融合、证据约束推理", "诊断解释与控制建议"],
        ["闭环控制层", "诊断建议、达标概率、回流收益、风险等级", "分层控制、规则约束、风险仲裁、必要时人工接管", "工程控制动作"],
    ],
    [1300, 2450, 3150, 2460],
)

section_break(doc)

add_heading(doc, "五、研究内容一：循环式水处理结构与时间缓冲机制", 1)
add_para(
    doc,
    "本研究内容解决“行动如何变得可行”的问题。低成本传感和复杂反应体系通常不支持瞬时判断，因此必须在工艺结构上提供可回流、可暂存、可延长停留和可二次调节的空间，使软传感器和多智能体诊断有足够时间形成可靠判断。",
)
add_table(
    doc,
    ["维度", "设计内容"],
    [
        ["目标", "设计一种可回流、可暂存、可延长停留、可切换处理路径的循环式水处理结构，把传感慢、反应慢、判断慢转化为可利用的时间窗口。"],
        ["关键问题", "循环什么时候有意义；回流几次才值得；延长停留与补加药剂哪个更优；何时停止循环并进入下一单元。"],
        ["技术方法", "构建主反应器、传感节点、回流缓冲池、加药单元、末端放行/切换单元；定义回流比、循环次数、停留时间、暂存时间窗；构建回流边际收益指标。"],
        ["实验设计", "设置一次通过、固定回流、动态回流三组对照；模拟进水浓度突增、pH 波动、氧化剂不足、催化剂衰减和传感延迟。"],
        ["评价指标", "去除率、达标概率、单位水量能耗、药剂消耗、处理时间、无效循环次数、风险出水拦截率。"],
        ["交付成果", "循环式小试/中试平台、循环控制参数库、回流边际收益评价模型、循环结构设计规范。"],
    ],
    [1500, 7860],
)
add_callout(
    doc,
    "本内容的核心贡献：",
    "把循环从传统工艺强化手段重新定义为时间缓冲、信息增益和风险拦截机制，使低成本传感条件下的智能控制具备工程可执行性。",
)

add_heading(doc, "六、研究内容二：低成本传感与软传感器灰箱状态估计", 1)
add_para(
    doc,
    "本研究内容解决“系统如何看见不可见状态”的问题。软传感器不应只预测出水浓度，而应估计中间过程状态，使系统能够判断异常来源、反应进展和控制需求。",
)
add_table(
    doc,
    ["状态类型", "建议估计变量", "控制意义"],
    [
        ["污染物状态", "目标污染物残留风险、达标概率、潜在副产物风险", "决定是否放行、回流或进入深度处理。"],
        ["反应状态", "反应完成度、氧化剂余量、催化剂活性、反应路径偏移", "判断是否延长停留、补加氧化剂或更换催化剂。"],
        ["基质状态", "高盐干扰、COD/TOC 负荷、pH 缓冲能力、传质限制", "判断是否需要预处理、调 pH 或切换工艺。"],
        ["传感状态", "传感器漂移、信号可信度、数据缺失风险", "决定是否暂停放行、触发旁路检测或人工接管。"],
        ["控制状态", "回流边际收益、继续处理成本、风险收益比", "决定回流、停留、加药和单元切换的优先级。"],
    ],
    [1500, 3620, 4240],
)
add_para(
    doc,
    "模型应采用“机理约束 + 时序学习 + 不确定性估计”的灰箱路线。输入不仅包括当前传感值，还应包括历史循环轨迹、投加记录、停留时间、回流次数、离线校准结果和故障标签。输出不仅是一个数值，而应包含状态估计、置信度、达标概率和建议验证动作。",
)
add_table(
    doc,
    ["模型组", "方法特征", "验证目的"],
    [
        ["纯黑箱 ML", "以低成本传感信号直接预测出水浓度", "作为最基础预测基线。"],
        ["传统软传感器", "估计目标水质指标，但缺少状态向量和机理约束", "验证软测量相对直接传感的价值。"],
        ["纯机理模型", "基于反应动力学、停留时间和药剂消耗建模", "验证机理先验的解释力和局限。"],
        ["灰箱软传感器", "融合机理约束、循环历史和不确定性估计", "验证本项目核心方法的优势。"],
    ],
    [1600, 4200, 3560],
)

add_heading(doc, "七、研究内容三：多智能体机理解释与故障诊断系统", 1)
add_para(
    doc,
    "本研究内容解决“系统如何解释异常并形成可信建议”的问题。受 ECOMATS 中知识图谱、多 agent 分工、多维评价和一致性融合思路启发，本项目将多智能体范式从材料发现迁移到水处理过程诊断与控制。",
)
add_table(
    doc,
    ["Agent", "职责", "输出"],
    [
        ["数据质控 Agent", "判断传感器漂移、数据缺失、异常尖峰和信号可信度", "数据质量标签、传感器健康度"],
        ["软测量 Agent", "调用软传感器估计隐藏状态并给出不确定性", "状态向量、置信度、达标概率"],
        ["机理解释 Agent", "结合知识图谱解释可能反应路径和抑制因素", "机理假设、证据链"],
        ["故障诊断 Agent", "识别药剂不足、pH 抑制、催化剂失活、基质冲击、传质限制等故障", "故障排序、原因置信度"],
        ["控制策略 Agent", "提出回流、延长停留、加药、调 pH、切换单元等动作", "候选控制策略"],
        ["成本安全 Agent", "评价能耗、药耗、时间成本、环境风险和误放行风险", "风险成本评分"],
        ["仲裁 Agent", "融合各 agent 结果，执行安全约束和一致性评分", "最终控制决策和诊断报告"],
    ],
    [1550, 4900, 2910],
)
add_para(
    doc,
    "知识图谱建议包含以下节点：污染物、水质基质、处理工艺、催化剂、氧化剂、传感信号、反应机理、故障模式、控制动作和证据来源。关系包括 degraded_by、activates、inhibits、indicates、caused_by、mitigates、has_cost、has_risk 和 supports。",
)
add_callout(
    doc,
    "多智能体约束原则：",
    "无证据不解释，高风险不放行，控制建议必须可执行，诊断报告必须能追溯到传感数据、软测量状态、知识图谱证据或历史实验记录。",
)

add_heading(doc, "八、研究内容四：循环条件下的闭环控制与工程验证", 1)
add_para(
    doc,
    "本研究内容解决“系统如何把判断转成动作”的问题。闭环控制不是附加模块，而是本方案从科研想法走向工程系统的关键。软传感器和多智能体输出必须最终落到泵、阀、加药、pH 调节、回流和单元切换上。",
)
add_table(
    doc,
    ["控制动作", "触发条件", "预期作用"],
    [
        ["达标放行", "达标概率高、状态置信度高、风险约束满足", "降低不必要处理时间和能耗。"],
        ["回流处理", "残留风险较高且回流边际收益为正", "增加反应机会并积累更多传感证据。"],
        ["延长停留", "反应未完成但药剂/催化剂仍有作用空间", "为慢反应和慢传感争取时间。"],
        ["补加药剂", "氧化剂不足、污染负荷突增或反应速率下降", "恢复反应驱动力。"],
        ["调节 pH", "pH 偏离最佳反应区间或出现抑制效应", "恢复反应条件和催化活性。"],
        ["切换单元", "循环收益低、基质干扰强或当前工艺失效", "进入吸附、膜、混凝、深度氧化等后续处理。"],
        ["暂存待检", "传感器异常、模型不确定性高或风险不可判定", "防止高风险出水误放行。"],
        ["报警接管", "连续诊断冲突、故障严重或安全约束触发", "引入人工判断和设备维护。"],
    ],
    [1600, 4500, 3260],
)
add_para(
    doc,
    "控制目标函数需要同时考虑达标概率、处理成本、能耗、时间、安全风险和回流边际收益。工程上可采用分层控制：底层控制泵阀和投加，中层根据软传感状态进行回流与停留控制，上层由多智能体系统完成诊断、策略选择和风险仲裁。",
)

section_break(doc)

add_heading(doc, "九、实验设计与验证体系", 1)
add_heading(doc, "9.1 场景包设计", 2)
add_table(
    doc,
    ["场景", "重点污染/过程", "验证重点"],
    [
        ["模拟有机污染废水", "染料、酚类、抗生素等可控体系", "验证平台、软传感器和基本闭环控制。"],
        ["高盐高 COD 废水", "复杂基质、氧化剂消耗、信号干扰", "验证基质干扰诊断和回流收益判断。"],
        ["含氟或芯片相关废水", "氟离子、pH 缓冲、复杂络合", "验证低成本传感下的风险拦截和达标概率判断。"],
        ["PFAS/PFOA 或替代体系", "持久污染物、高级氧化、潜在副产物", "验证机理解释、深度处理和安全放行策略。"],
    ],
    [1900, 3150, 4310],
)

add_heading(doc, "9.2 对照组设计", 2)
add_table(
    doc,
    ["组别", "系统配置", "证明目的"],
    [
        ["A", "一次通过，无 AI", "建立传统基线。"],
        ["B", "固定回流，无软传感", "证明单纯循环不能解决诊断与优化问题。"],
        ["C", "黑箱 ML 预测出水", "证明预测不等于解释和控制。"],
        ["D", "软传感器 + 人工决策", "证明状态估计的价值。"],
        ["E", "软传感器 + 多智能体诊断", "证明解释和故障识别能力。"],
        ["F", "完整闭环系统", "证明整体方案在达标、成本、风险和自动化上的综合优势。"],
    ],
    [1000, 3450, 4910],
)

add_heading(doc, "9.3 扰动与故障库", 2)
add_table(
    doc,
    ["扰动/故障", "系统应识别的状态", "期望控制动作"],
    [
        ["进水浓度突增", "污染物残留风险上升、达标概率下降", "回流、延长停留、补加药剂或暂存。"],
        ["pH 偏离", "反应抑制、催化活性下降", "调 pH 后继续处理。"],
        ["氧化剂不足", "反应完成度不足、氧化能力下降", "补加氧化剂并观察回流后状态。"],
        ["催化剂衰减", "催化活性持续下降、回流收益降低", "再生、更换、提高停留或切换单元。"],
        ["高盐/高 COD 干扰", "基质竞争、氧化剂无效消耗", "预处理、提高剂量、切换工艺或停止无效循环。"],
        ["传感器漂移", "信号可信度下降、诊断冲突增加", "暂存待检、触发校准或旁路检测。"],
    ],
    [2200, 3550, 3610],
)

add_heading(doc, "9.4 评价指标", 2)
add_label_block(doc, "水质效果：", "污染物去除率、达标率、达标概率预测准确率、高风险出水误放行率。")
add_label_block(doc, "感知效果：", "隐藏状态估计误差、异常提前识别时间、传感器漂移鲁棒性、离线真值校准效率。")
add_label_block(doc, "诊断效果：", "故障诊断准确率、误诊率、漏诊率、诊断解释可信度、与专家判断一致性。")
add_label_block(doc, "控制效果：", "无效回流减少比例、药剂节约率、能耗变化、处理时间变化、人工干预次数。")
add_label_block(doc, "工程效果：", "系统连续运行稳定性、控制动作可执行率、维护复杂度、可迁移场景数量。")

add_heading(doc, "十、年度任务与里程碑", 1)
add_table(
    doc,
    ["年度", "核心任务", "阶段成果", "验收指标"],
    [
        ["第 1 年", "搭建循环式水处理平台，完成低成本传感体系、数据采集系统和标准实验流程。", "小试平台、传感配置、初始数据集、实验 SOP。", "平台可连续运行；传感数据完整率大于 90%；实现回流、加药、延长停留等基础动作。"],
        ["第 2 年", "建立软传感器灰箱估计模型，形成隐藏状态向量、不确定性估计和传感器健康诊断。", "软传感器模型库、隐藏状态数据集、传感器健康模块。", "关键状态估计优于黑箱基线；达标概率预测准确率提高；异常提前识别能力形成。"],
        ["第 3 年", "构建知识图谱与多智能体诊断系统，完成故障模式库和诊断报告模块。", "知识图谱、多 agent 原型、故障库、诊断报告生成工具。", "诊断准确率优于单一模型；控制建议与专家判断一致性提高；高风险错误建议下降。"],
        ["第 4 年", "集成闭环控制系统，开展小试/中试和真实废水场景验证。", "闭环控制原型、中试验证平台、示范报告、论文专利软件成果。", "达标率提升；风险出水误放行下降；药剂或能耗降低；人工干预次数减少。"],
    ],
    [900, 3250, 2500, 2710],
)

add_heading(doc, "十一、平台配置与团队分工", 1)
add_heading(doc, "11.1 平台配置", 2)
add_table(
    doc,
    ["平台", "组成", "用途"],
    [
        ["实验平台", "反应器、进水池、回流缓冲池、出水池、泵阀、加药单元、pH 调节单元、旁路采样口", "形成可回流、可暂存、可调节的水处理物理系统。"],
        ["传感平台", "pH、ORP、电导率、浊度、温度、流量、液位、UV254，必要时增加氟离子电极和快检模块", "构建低成本过程证据链。"],
        ["数据平台", "边缘采集设备、数据库、实验日志、离线检测导入、数据清洗与标注工具", "支撑软传感器训练和多智能体诊断。"],
        ["AI 平台", "软传感器模型、知识图谱、多 agent 调度、诊断报告、控制策略库、风险仲裁模块", "完成灰箱估计、机理解释和策略生成。"],
        ["控制平台", "PLC 或边缘控制器、泵阀执行系统、投加控制、回流控制、报警和人工接管界面", "把 AI 判断转化为实际工程动作。"],
    ],
    [1350, 5110, 2900],
)

add_heading(doc, "11.2 团队分工", 2)
add_table(
    doc,
    ["子课题/团队", "主要职责", "关键交付"],
    [
        ["工艺与反应器团队", "循环结构、回流路径、反应器运行、污染场景设计", "平台、工艺参数库、实验数据。"],
        ["传感与数据团队", "低成本传感器布设、校准、漂移识别、数据采集", "传感配置、数据集、数据质量模型。"],
        ["软传感建模团队", "隐藏状态定义、灰箱模型、时序学习、不确定性估计", "软传感器模型库、状态估计模块。"],
        ["多智能体与知识工程团队", "知识图谱、agent 分工、诊断推理、报告生成", "多 agent 系统、故障库、诊断报告。"],
        ["控制工程团队", "闭环控制策略、泵阀执行、风险仲裁、系统集成", "控制原型、执行策略库、中试验证。"],
    ],
    [1700, 4550, 3110],
)

add_heading(doc, "十二、预期创新点", 1)
add_label_block(doc, "创新点一：", "提出“循环停留驱动的信息增益”概念，把回流从单纯工艺强化手段升级为低成本传感条件下的时间缓冲和风险拦截机制。")
add_label_block(doc, "创新点二：", "构建面向过程隐藏状态的软传感器，不只预测出水浓度，而是估计污染物残留、反应完成度、氧化剂余量、催化剂活性和基质干扰等灰箱状态。")
add_label_block(doc, "创新点三：", "将多智能体科研发现范式迁移到水处理过程诊断与控制，形成数据质控、机理解释、故障诊断、成本安全评价和仲裁决策的协同框架。")
add_label_block(doc, "创新点四：", "建立回流、延时、加药、调 pH、切换单元和风险放行之间的闭环决策机制，使 AI 结果真正进入工程控制动作。")
add_label_block(doc, "创新点五：", "形成低成本传感条件下从黑箱到灰箱再到闭环控制的系统性方法，为复杂废水智能治理提供可复制范式。")

add_heading(doc, "十三、风险与应对策略", 1)
add_table(
    doc,
    ["风险", "表现", "应对策略"],
    [
        ["软传感器估计不稳定", "少样本、传感噪声、场景迁移导致误差大", "引入离线真值校准、机理约束、不确定性输出和主动采样。"],
        ["多智能体诊断幻觉", "给出缺少证据的解释或高风险建议", "要求证据引用、一致性融合、安全规则硬约束和人工接管机制。"],
        ["循环导致成本上升", "过度回流、处理时间过长、能耗增加", "建立回流边际收益和最大循环次数约束。"],
        ["真实废水波动过大", "模型从模拟水迁移到真实水失败", "采用分阶段场景递进、在线校准和知识图谱更新。"],
        ["控制动作不可执行", "AI 建议超出设备能力或工程边界", "建立控制动作白名单、设备约束表和执行前安全仲裁。"],
    ],
    [1900, 3480, 3980],
)

add_heading(doc, "十四、预期成果", 1)
add_table(
    doc,
    ["成果类型", "具体成果"],
    [
        ["理论方法", "循环停留信息增益理论、低成本传感灰箱状态估计方法、多智能体诊断仲裁方法。"],
        ["模型算法", "软传感器模型库、传感器健康诊断模型、回流边际收益模型、闭环控制策略库。"],
        ["平台系统", "循环式水处理小试/中试平台、多智能体诊断系统、闭环控制软件原型。"],
        ["数据资源", "典型污染场景数据集、故障扰动数据集、传感-离线真值配对数据集。"],
        ["论文专利", "软传感灰箱估计、多智能体诊断、循环闭环控制、低成本智能水处理系统等方向论文与专利。"],
        ["工程示范", "面向含氟废水、高盐高 COD 废水或 PFAS/PFOA 场景的验证报告和应用边界说明。"],
    ],
    [1650, 7710],
)

add_heading(doc, "十五、参考线索与文献拓展方向", 1)
add_para(
    doc,
    "前期讨论的重要触发点来自 Nature Water 2026 年论文“Multi-agent artificial intelligence designs novel catalysts for ultrafast water purification”（DOI: 10.1038/s44221-026-00634-9）。该研究提出 ECOMATS，多智能体系统整合专家验证知识图谱、多个微调大模型、五维评价和多 agent 盲评融合，用于高级氧化催化剂发现。本文档不是照搬其材料发现任务，而是借鉴其“知识图谱 + 多智能体分工 + 多维评价 + 一致性融合”的科研组织范式，并将对象迁移到水处理过程诊断和闭环控制。",
)
add_label_block(doc, "建议检索方向：", "low-cost sensing wastewater treatment；soft sensor water treatment process；grey-box model wastewater treatment；sparse sensing process control；recirculation hydraulic retention time wastewater；multi-agent system water treatment；closed-loop control advanced oxidation process；fault diagnosis wastewater treatment；physics-informed machine learning water treatment。")

add_heading(doc, "十六、最终凝练", 1)
add_callout(
    doc,
    "一句话版本：",
    "本项目试图在传感器买不起、过程看不清、反应来不及判断的真实工程条件下，通过循环结构把时间变成信息，再用软传感器和多智能体系统把信息变成可解释、可执行、可闭环的水处理控制行动。",
)
add_para(
    doc,
    "它的真正价值不在于简单把 AI 用到水处理，而在于提出一套面向低成本约束的工程智能化路径：先承认过程黑箱，再用循环创造观察机会，用软传感器恢复关键状态，用多智能体解释原因和权衡风险，最后用闭环控制把判断落实为具体动作。这一思路既可以作为重大项目研究方案，也可以拆分为博士/硕士课题、论文体系、平台建设任务和工程示范路线。",
)

doc.save(OUT)
print(OUT)
